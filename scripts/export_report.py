#!/usr/bin/env python3
"""Export SurveilAMR report to DOCX and PDF with embedded figures, footers, and page numbers."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "SurveilAMR_Report.md"
DOCX_PATH = ROOT / "docs" / "SurveilAMR_Report.docx"
PDF_PATH = ROOT / "docs" / "SurveilAMR_Report.pdf"

FOOTER_TEXT = "2026 Vivli AMR Surveillance Data Challenge"
MAIN_BODY_MAX_PAGES = 5
IMG_WIDTH = Inches(4.2)
BODY_FONT = Pt(9.5)
IMG_HEIGHT_DOCX = Inches(2.0)


def parse_md(path: Path) -> tuple[list, list]:
    """Return (main_blocks, reference_blocks) split at ## References."""
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks: list = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# "):
            blocks.append(("h0", line[2:].strip()))
        elif line.startswith("## "):
            blocks.append(("h1", line[3:].strip()))
        elif line.startswith("### "):
            blocks.append(("h2", line[4:].strip()))
        elif line.startswith("!["):
            m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line)
            if m:
                rel = m.group(2).lstrip("../").lstrip("/")
                img_path = ROOT / rel if not Path(m.group(2)).is_absolute() else Path(m.group(2))
                if not img_path.exists():
                    img_path = ROOT / m.group(2).replace("../", "")
                blocks.append(("img", m.group(1), img_path))
        elif line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|---"):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            blocks.append(("table", table_lines))
            continue
        elif line.startswith(">"):
            blocks.append(("quote", line.lstrip("> ").strip()))
        elif line.strip() == "---":
            pass
        elif line.strip():
            blocks.append(("p", line.strip()))
        i += 1

    ref_idx = next((idx for idx, b in enumerate(blocks) if b[0] == "h1" and b[1] == "References"), None)
    if ref_idx is None:
        return blocks, []
    return blocks[:ref_idx], blocks[ref_idx:]


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)

    run4 = paragraph.add_run("1")
    run5 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)


def _set_footer(section, centered: bool = True) -> None:
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(FOOTER_TEXT)
    p.add_run("  |  Page ")
    _add_page_number(p)


def _apply_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = BODY_FONT
    for level in range(1, 4):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(14 - level)


def _render_blocks(doc: Document, blocks: list) -> None:
    for block in blocks:
        kind = block[0]
        if kind == "h0":
            p = doc.add_heading(block[1], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "h1":
            doc.add_heading(block[1], level=1)
        elif kind == "h2":
            doc.add_heading(block[1], level=2)
        elif kind == "img":
            caption, img_path = block[1], block[2]
            if img_path.exists():
                doc.add_picture(str(img_path), width=IMG_WIDTH, height=IMG_HEIGHT_DOCX)
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)
        elif kind == "table":
            rows = [[c.strip() for c in tl.strip("|").split("|")] for tl in block[1]]
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for r, row in enumerate(rows):
                    for c, cell in enumerate(row):
                        table.rows[r].cells[c].text = cell
                        for para in table.rows[r].cells[c].paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(9)
        elif kind == "quote":
            doc.add_paragraph(block[1])
        elif kind == "p":
            para = doc.add_paragraph(block[1])
            para.paragraph_format.space_after = Pt(4)


def build_docx(main_blocks: list, ref_blocks: list) -> None:
    doc = Document()
    _apply_styles(doc)

    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    _render_blocks(doc, main_blocks)

    if ref_blocks:
        doc.add_page_break()
        new_section = doc.add_section()
        new_section.top_margin = Inches(0.85)
        new_section.bottom_margin = Inches(0.9)
        new_section.left_margin = Inches(0.9)
        new_section.right_margin = Inches(0.9)
        _render_blocks(doc, ref_blocks)

    for section in doc.sections:
        _set_footer(section)

    doc.save(DOCX_PATH)
    print(f"Saved {DOCX_PATH}")


def docx_to_pdf() -> None:
    """Export PDF via Word COM (Windows) with footer preserved."""
    try:
        import win32com.client  # type: ignore

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(DOCX_PATH.resolve()))

        for section in doc.Sections:
            footer = section.Footers(1)
            footer.Range.Text = ""
            footer.Range.ParagraphFormat.Alignment = 1  # wdAlignParagraphCenter
            footer.Range.Text = f"{FOOTER_TEXT}  |  Page "
            footer.PageNumbers.Add(1)  # wdAlignPageNumberCenter

        doc.SaveAs(str(PDF_PATH.resolve()), FileFormat=17)
        page_count = doc.ComputeStatistics(2)  # wdStatisticPages
        ref_start = None
        for i in range(1, doc.Sections.Count + 1):
            if doc.Sections(i).Range.Text.strip().startswith("References"):
                ref_start = doc.Sections(i).Range.Information(3)  # wdActiveEndPageNumber
                break
        main_pages = (ref_start - 1) if ref_start and ref_start > 1 else page_count
        doc.Close(False)
        word.Quit()
        print(f"Saved {PDF_PATH} (via Word, {page_count} total pages)")
        print(f"Main body: ~{main_pages} pages (References excluded from 5-page limit; target <= {MAIN_BODY_MAX_PAGES})")
        return
    except Exception as exc:
        print(f"Word PDF export unavailable ({exc}); using reportlab fallback")
        _pdf_reportlab_fallback()


def _pdf_reportlab_fallback() -> None:
    """Build PDF with embedded figures and footers when Word is unavailable."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas

    main_blocks, ref_blocks = parse_md(MD_PATH)
    c = canvas.Canvas(str(PDF_PATH), pagesize=letter)
    width, height = letter
    margin = 0.75 * inch
    y = height - margin
    page_num = 0
    in_refs = False

    def new_page(start_refs: bool = False) -> None:
        nonlocal y, page_num, in_refs
        if page_num > 0:
            _draw_footer(c, page_num, width)
            c.showPage()
        page_num += 1
        in_refs = start_refs
        y = height - margin

    def _draw_footer(cv, num: int, w: float) -> None:
        cv.setFont("Helvetica", 8)
        cv.drawCentredString(w / 2, 0.45 * inch, f"{FOOTER_TEXT}  |  Page {num}")

    new_page()

    for block in main_blocks + ref_blocks:
        if block[0] == "h1" and block[1] == "References" and not in_refs:
            new_page(start_refs=True)

        if block[0] == "h0":
            if y < margin + 60:
                new_page(in_refs)
            c.setFont("Helvetica-Bold", 14)
            c.drawCentredString(width / 2, y, block[1][:90])
            y -= 22
        elif block[0] == "h1":
            if y < margin + 40:
                new_page(in_refs)
            c.setFont("Helvetica-Bold", 12)
            y -= 8
            c.drawString(margin, y, block[1])
            y -= 18
        elif block[0] == "h2":
            if y < margin + 30:
                new_page(in_refs)
            c.setFont("Helvetica-Bold", 10.5)
            c.drawString(margin, y, block[1])
            y -= 16
        elif block[0] == "img":
            caption, img_path = block[1], block[2]
            if img_path.exists():
                img_h = 3.0 * inch
                if y - img_h < margin + 30:
                    new_page(in_refs)
                c.drawImage(ImageReader(str(img_path)), margin, y - img_h, width=5.4 * inch, height=img_h, preserveAspectRatio=True, anchor="nw")
                y -= img_h + 4
                c.setFont("Helvetica-Oblique", 8)
                for line in _wrap(caption, 85):
                    c.drawCentredString(width / 2, y, line)
                    y -= 10
                y -= 6
        elif block[0] == "p":
            c.setFont("Helvetica", 9.5)
            for line in _wrap(block[1], 95):
                if y < margin + 14:
                    new_page(in_refs)
                c.drawString(margin, y, line)
                y -= 12
        elif block[0] == "quote":
            c.setFont("Helvetica-Oblique", 9)
            for line in _wrap(block[1], 90):
                if y < margin + 14:
                    new_page(in_refs)
                c.drawString(margin + 10, y, line)
                y -= 12

    if page_num > 0:
        _draw_footer(c, page_num, width)
    c.save()
    print(f"Saved {PDF_PATH} (reportlab fallback)")


def _wrap(text: str, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current: list[str] = []
    for word in words:
        trial = " ".join(current + [word])
        if len(trial) <= width:
            current.append(word)
        else:
            if current:
                lines.append(" ".join(current))
            current = [word]
    if current:
        lines.append(" ".join(current))
    return lines or [""]


if __name__ == "__main__":
    main_blocks, ref_blocks = parse_md(MD_PATH)
    build_docx(main_blocks, ref_blocks)
    docx_to_pdf()
